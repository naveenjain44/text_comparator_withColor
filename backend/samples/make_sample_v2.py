"""Sample with rich HTML paragraphs (like a real marketing email) + strikethrough + placement swap."""
import os
from email.message import EmailMessage
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/app/backend/samples"
os.makedirs(OUT, exist_ok=True)

# ----- DOCX mockup with strikethrough runs and 10-ish real paragraphs -----
d = Document()
d.add_paragraph("Subject: We're updating our Terms of Service")
d.add_paragraph("Hi there,")

d.add_paragraph(
    "Every couple of years, we update our Terms of Service. We wanted to let you know ahead "
    "of time that the next update will be on July 30, 2026."
)

# strikethrough draft text within a real paragraph
p = d.add_paragraph()
p.add_run(
    "These changes won't affect the way you use our services, but they should "
)
strike_run = p.add_run("[DRAFT LEGAL NOTE — DELETE BEFORE SEND] ")
strike_run.font.strike = True
p.add_run(
    "help make it easier for you to understand what to expect from Google — and "
    "what we expect from you — as you use our services."
)

d.add_paragraph(
    "You can review the new terms here. If you're based in the European Economic Area (EEA) "
    "or Switzerland, we've also provided a summary of the key changes to the EEA version of our terms."
)
d.add_paragraph("At a glance, here's what this update means for you:")
d.add_paragraph("In general:", style="List Bullet")
d.add_paragraph(
    "We added a section to help you better understand why our services may access the internet "
    "when not actively engaged, to encourage you to check your Internet service plan and your device "
    "and network settings, as each of those may affect your costs."
)
d.add_paragraph(
    "We updated and clarified our Settling disputes, governing law, and courts section."
)
d.add_paragraph(
    "We make it clearer how various sections in our terms relate to each other."
)
d.add_paragraph(
    "Thank you for using Google services!"
)

# CTA
p = d.add_paragraph()
part = d.part
rid = part.relate_to(
    "https://policies.google.com/terms",
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
    is_external=True,
)
hyperlink = OxmlElement("w:hyperlink")
hyperlink.set(qn("r:id"), rid)
r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr"); r.append(rPr)
t = OxmlElement("w:t"); t.text = "Review the new terms"; r.append(t)
hyperlink.append(r)
p._p.append(hyperlink)

d.add_paragraph("Unsubscribe | Privacy Policy | © 2026 Google LLC · 1600 Amphitheatre Pkwy")
d.save(f"{OUT}/tos_v2.docx")

# ----- EML output: same paragraphs but "Thank you..." moved BEFORE "We make it clearer..." (placement swap) -----
msg = EmailMessage()
msg["Subject"] = "We're updating our Terms of Service"
msg["From"] = "noreply@google.com"
msg["To"] = "user@example.com"
html = """
<html><body>
  <p>Hi there,</p>
  <p>Every couple of years, we update our Terms of Service. We wanted to let you know ahead of time that the next update will be on July 30, 2026.</p>
  <p>These changes won't affect the way you use our services, but they should help make it easier for you to understand what to expect from Google &mdash; and what we expect from you &mdash; as you use our services.</p>
  <p>You can review the new terms here. If you're based in the European Economic Area (EEA) or Switzerland, we've also provided a summary of the key changes to the EEA version of our terms.</p>
  <p>At a glance, here's what this update means for you:</p>
  <ul><li>In general:</li></ul>
  <p>We added a section to help you better understand why our services may access the internet when not actively engaged, to encourage you to check your Internet service plan and your device and network settings, as each of those may affect your costs.</p>
  <p>We updated and clarified our Settling disputes, governing law, and courts section.</p>
  <p>Thank you for using Google services!</p>
  <p>We make it clearer how various sections in our terms relate to each other.</p>
  <p><a href="https://policies.google.com/terms">Review the new terms</a></p>
  <p><img src="https://cdn.google.com/logo.png" alt="Google Logo"/></p>
  <p>Unsubscribe | Privacy Policy | &copy; 2026 Google LLC &middot; 1600 Amphitheatre Pkwy</p>
</body></html>
"""
msg.set_content("plain fallback")
msg.add_alternative(html, subtype="html")
with open(f"{OUT}/tos_v2.eml", "wb") as f:
    f.write(bytes(msg))

print("Wrote:", os.listdir(OUT))
