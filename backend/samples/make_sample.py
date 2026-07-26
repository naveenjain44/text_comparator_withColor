"""Create sample DOCX (with strikethrough) + EML to validate scoring & strikethrough exclusion."""
import io, os
from email.message import EmailMessage
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

os.makedirs("/tmp/uat_samples", exist_ok=True)

# --- DOCX mockup with strikethrough ---
d = Document()
d.add_paragraph("Subject: Terms of Service update — action needed")
d.add_paragraph("Hi there,")

d.add_paragraph(
    "Every couple of years, we update our Terms of Service. We wanted to let you know ahead of "
    "time that the next update will be on July 30, 2026."
)

# paragraph WITH strikethrough content that must NOT be compared
p = d.add_paragraph()
r1 = p.add_run("These changes won't affect the way you use our services, but they should ")
r1.font.strike = False
r2 = p.add_run("[OLD LEGAL DRAFT TEXT WE MUST REMOVE] ")
r2.font.strike = True  # strikethrough — should be excluded
r3 = p.add_run("help make it easier for you to understand what to expect from Google — and what we expect from you — as you use our services.")

d.add_paragraph(
    "You can review the new terms here. If you're based in the European Economic Area (EEA) or "
    "Switzerland, we've also provided a summary of the key changes to the EEA version of our terms."
)
d.add_paragraph("At a glance, here's what this update means for you:")

# CTA
p = d.add_paragraph()
part = d.part
rid = part.relate_to(
    "https://policies.google.com/terms",
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
    is_external=True,
)
hyperlink = OxmlElement("w:hyperlink")
hyperlink.set(qn("r:id"), rid)
new_run = OxmlElement("w:r")
rPr = OxmlElement("w:rPr")
new_run.append(rPr)
t = OxmlElement("w:t")
t.text = "Review the new terms"
new_run.append(t)
hyperlink.append(new_run)
p._p.append(hyperlink)

d.add_paragraph("Unsubscribe | Privacy Policy | © 2026 Google LLC · 1600 Amphitheatre Pkwy")
d.save("/tmp/uat_samples/tos_update.docx")

# --- EML output (matches DOCX minus strikethrough) ---
msg = EmailMessage()
msg["Subject"] = "Terms of Service update — action needed"
msg["From"] = "noreply@google.com"
msg["To"] = "user@example.com"
html = """
<html><body>
<p>Hi there,</p>
<p>Every couple of years, we update our Terms of Service. We wanted to let you know ahead of time that the next update will be on July 30, 2026.</p>
<p>These changes won't affect the way you use our services, but they should help make it easier for you to understand what to expect from Google &mdash; and what we expect from you &mdash; as you use our services.</p>
<p>You can review the new terms here. If you're based in the European Economic Area (EEA) or Switzerland, we've also provided a summary of the key changes to the EEA version of our terms.</p>
<p>At a glance, here's what this update means for you:</p>
<p><a href="https://policies.google.com/terms">Review the new terms</a></p>
<p><img src="https://cdn.google.com/logo.png" alt="Google Logo"/></p>
<p>Unsubscribe | Privacy Policy | &copy; 2026 Google LLC &middot; 1600 Amphitheatre Pkwy</p>
</body></html>
"""
msg.set_content("Plain fallback")
msg.add_alternative(html, subtype="html")
with open("/tmp/uat_samples/tos_update.eml", "wb") as f:
    f.write(bytes(msg))

print("wrote", os.listdir("/tmp/uat_samples"))
