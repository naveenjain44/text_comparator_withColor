"""
Core parser + comparator for DOCX mockup vs EML/MSG output.
100% offline. No network calls at runtime.

Improvements in this version
----------------------------
* Strikethrough (`w:strike` / `w:dstrike`) text is **excluded** from the mockup
  (both plain runs and hyperlink runs) — matches the user's requirement.
* Section scores are computed as **percentages** and combined into a weighted
  overall score for accurate reporting.
* STRICT / SMART modes:
    - STRICT: exact match after whitespace normalization; no glossary.
    - SMART:  glossary-aware, case-insensitive, punctuation-tolerant.
"""
from __future__ import annotations

import email
import io
import re
from dataclasses import dataclass, field, asdict
from difflib import SequenceMatcher
from email import policy
from typing import List, Dict, Optional, Tuple

from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn

try:
    import extract_msg
    HAS_MSG = True
except Exception:
    HAS_MSG = False


# ---------- Data models ----------

@dataclass
class LinkItem:
    text: str
    url: str


@dataclass
class ImageItem:
    filename: str
    alt: str


@dataclass
class ParsedEmail:
    subject: str = ""
    greeting: str = ""
    body_paragraphs: List[str] = field(default_factory=list)
    cta: Optional[LinkItem] = None
    footer: str = ""
    images: List[ImageItem] = field(default_factory=list)
    links: List[LinkItem] = field(default_factory=list)
    raw_text: str = ""

    def to_dict(self):
        return {
            "subject": self.subject,
            "greeting": self.greeting,
            "body_paragraphs": self.body_paragraphs,
            "cta": asdict(self.cta) if self.cta else None,
            "footer": self.footer,
            "images": [asdict(i) for i in self.images],
            "links": [asdict(l) for l in self.links],
            "raw_text": self.raw_text,
        }


# ---------- Constants ----------

CTA_KEYWORDS = [
    "shop now", "buy now", "get started", "learn more", "sign up",
    "register", "book now", "start", "join", "download", "explore",
    "discover", "view", "claim", "redeem", "activate", "continue",
    "get it", "try", "order now", "read more", "see more", "review the new terms",
]

GREETING_PATTERNS = [
    r"^\s*(hi|hello|hey|dear|greetings|welcome|good\s+(morning|afternoon|evening))\b[^\n]*",
]

FOOTER_KEYWORDS = [
    "unsubscribe", "privacy policy", "terms of service", "terms & conditions",
    "copyright", "©", "all rights reserved", "you are receiving",
    "manage preferences", "contact us", "follow us", "view online",
    "google llc", "1600 amphitheatre",
]


# ---------- Text utilities ----------

_PUNCT_RE = re.compile(r"[^\w\s]", re.UNICODE)


def normalize_ws(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())


def _smart_norm(text: str) -> str:
    t = normalize_ws(text).lower()
    t = _PUNCT_RE.sub(" ", t)
    return normalize_ws(t)


def apply_glossary(text: str, glossary: List[Dict[str, str]]) -> str:
    if not text or not glossary:
        return text or ""
    out = text
    for rule in glossary:
        orig = (rule.get("original") or "").strip()
        variation = (rule.get("variation") or "").strip()
        if not orig or not variation:
            continue
        out = re.compile(re.escape(variation), re.IGNORECASE).sub(orig, out)
    return out


def norm_for_compare(text: str, glossary, mode: str = "smart") -> str:
    if mode == "strict":
        return normalize_ws(text or "")
    return _smart_norm(apply_glossary(text or "", glossary))


def similarity(a: str, b: str, glossary, mode: str) -> float:
    an, bn = norm_for_compare(a, glossary, mode), norm_for_compare(b, glossary, mode)
    if not an and not bn:
        return 1.0
    if not an or not bn:
        return 0.0
    return SequenceMatcher(None, an, bn).ratio()


def is_greeting(text: str) -> bool:
    t = (text or "").strip().lower()
    return any(re.match(p, t) for p in GREETING_PATTERNS)


def is_footer_line(text: str) -> bool:
    t = (text or "").lower()
    return any(k in t for k in FOOTER_KEYWORDS)


def is_cta_text(text: str) -> bool:
    t = (text or "").strip().lower()
    if not t or len(t) > 60:
        return False
    return any(k in t for k in CTA_KEYWORDS)


# ---------- DOCX helpers (strikethrough-aware) ----------

def _run_is_strike(r_elem) -> bool:
    """Return True if <w:r> element has strikethrough formatting."""
    for child in r_elem:
        if child.tag.endswith("}rPr"):
            for sub in child:
                tag = sub.tag.split("}")[-1]
                if tag in ("strike", "dstrike"):
                    val = sub.get(qn("w:val"))
                    if val is None:
                        return True
                    if str(val).lower() in ("true", "1", "on"):
                        return True
            return False
    return False


def _text_from_run(r_elem) -> str:
    if _run_is_strike(r_elem):
        return ""
    parts = []
    for t in r_elem.iter():
        tag = t.tag.split("}")[-1]
        if tag == "t":
            parts.append(t.text or "")
        elif tag == "tab":
            parts.append(" ")
        elif tag == "br":
            parts.append("\n")
    return "".join(parts)


def _find_docPr_alt(anchor) -> str:
    """Locate the docPr descr/title alt text around a drawing anchor."""
    parent = anchor.getparent()
    hops = 0
    while parent is not None and hops < 6:
        for sub in parent.iter():
            if sub.tag.endswith("}docPr"):
                return (sub.get("descr") or sub.get("title") or "").strip()
        parent = parent.getparent()
        hops += 1
    return ""


def _extract_paragraph(p_elem, rels: Dict[str, str], image_rels: Dict[str, str]):
    """
    Walk children of <w:p> in order and return:
      text (str, strikethrough excluded)
      links (list[LinkItem])
      images (list[ImageItem])
    """
    text_parts: List[str] = []
    links: List[LinkItem] = []
    images: List[ImageItem] = []

    def _walk_run(r_elem):
        # image blips inside this run
        for blip in r_elem.iter():
            if blip.tag.endswith("}blip"):
                embed = blip.get(qn("r:embed"))
                fname = image_rels.get(embed, embed or "image")
                alt = _find_docPr_alt(blip)
                images.append(ImageItem(filename=fname, alt=normalize_ws(alt)))
        return _text_from_run(r_elem)

    for child in p_elem:
        tag = child.tag.split("}")[-1]
        if tag == "r":
            text_parts.append(_walk_run(child))
        elif tag == "hyperlink":
            rid = child.get(qn("r:id"))
            url = rels.get(rid, "")
            link_text_parts = []
            for r in child:
                if r.tag.endswith("}r"):
                    link_text_parts.append(_walk_run(r))
            link_text = normalize_ws("".join(link_text_parts))
            if link_text or url:
                links.append(LinkItem(text=link_text, url=url))
                text_parts.append(link_text)

    return normalize_ws("".join(text_parts)), links, images


# ---------- DOCX parser ----------

def parse_docx(file_bytes: bytes) -> ParsedEmail:
    doc = Document(io.BytesIO(file_bytes))
    result = ParsedEmail()

    rels: Dict[str, str] = {}
    image_rels: Dict[str, str] = {}
    try:
        for rel_id, rel in doc.part.rels.items():
            if "hyperlink" in rel.reltype:
                rels[rel_id] = rel.target_ref
            elif "image" in rel.reltype:
                image_rels[rel_id] = rel.target_ref.split("/")[-1]
    except Exception:
        pass

    body_lines: List[str] = []
    for para in doc.paragraphs:
        text, links, images = _extract_paragraph(para._p, rels, image_rels)
        result.links.extend(links)
        result.images.extend(images)
        if text:
            body_lines.append(text)

    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text, links, images = _extract_paragraph(p._p, rels, image_rels)
                    result.links.extend(links)
                    result.images.extend(images)
                    if text and text not in body_lines:
                        body_lines.append(text)

    _split_sections(result, body_lines)
    result.raw_text = "\n".join(body_lines)
    return result


# ---------- EML / MSG parsers ----------

def parse_eml(file_bytes: bytes) -> ParsedEmail:
    msg = email.message_from_bytes(file_bytes, policy=policy.default)
    result = ParsedEmail()
    result.subject = normalize_ws(msg.get("Subject", "") or "")

    html_body = None
    text_body = None
    for part in msg.walk():
        ctype = part.get_content_type()
        cdisp = str(part.get("Content-Disposition") or "")
        if ctype == "text/html" and "attachment" not in cdisp and html_body is None:
            try:
                html_body = part.get_content()
            except Exception:
                payload = part.get_payload(decode=True)
                if payload:
                    html_body = payload.decode(part.get_content_charset() or "utf-8", errors="replace")
        elif ctype == "text/plain" and "attachment" not in cdisp and text_body is None:
            try:
                text_body = part.get_content()
            except Exception:
                payload = part.get_payload(decode=True)
                if payload:
                    text_body = payload.decode(part.get_content_charset() or "utf-8", errors="replace")
        elif ctype.startswith("image/"):
            fname = part.get_filename() or (part.get("Content-ID", "") or "image").strip("<>")
            # dedupe later after HTML parse
            result.images.append(ImageItem(filename=fname, alt=""))

    if html_body:
        _parse_html_body(html_body, result)
    elif text_body:
        _parse_plain_body(text_body, result)

    return result


def parse_msg(file_bytes: bytes) -> ParsedEmail:
    if not HAS_MSG:
        raise RuntimeError("extract-msg is not installed")
    with io.BytesIO(file_bytes) as f:
        m = extract_msg.Message(f)
        result = ParsedEmail()
        result.subject = normalize_ws(m.subject or "")
        html = m.htmlBody
        if isinstance(html, bytes):
            html = html.decode("utf-8", errors="replace")
        if html:
            _parse_html_body(html, result)
        elif m.body:
            _parse_plain_body(m.body, result)

        try:
            for att in m.attachments:
                name = getattr(att, "longFilename", None) or getattr(att, "shortFilename", None) or "attachment"
                if isinstance(name, str) and re.search(r"\.(png|jpe?g|gif|webp|svg|bmp)$", name, re.I):
                    if not any(i.filename == name for i in result.images):
                        result.images.append(ImageItem(filename=name, alt=""))
        except Exception:
            pass
        return result


def _parse_html_body(html: str, result: ParsedEmail):
    soup = BeautifulSoup(html, "lxml")

    # remove hidden preview text / trackers
    for tag in soup.find_all(True, attrs={"style": True}):
        style = (tag.get("style") or "").lower().replace(" ", "")
        if "display:none" in style or "max-height:0" in style or "font-size:0" in style:
            tag.decompose()
    for tag in soup.find_all(["script", "style"]):
        tag.decompose()

    # Images (preserve order in flow)
    seen_images = set()
    for img in soup.find_all("img"):
        src = img.get("src") or ""
        fname = src.split("?")[0].split("/")[-1] if src else (img.get("alt") or "image")
        fname = re.sub(r"^cid:", "", fname)
        alt = normalize_ws(img.get("alt") or "")
        key = (fname, alt)
        if key in seen_images:
            continue
        seen_images.add(key)
        result.images.append(ImageItem(filename=fname or "image", alt=alt))

    # Hyperlinks
    for a in soup.find_all("a"):
        href = a.get("href") or ""
        text = normalize_ws(a.get_text(" ", strip=True))
        if href.startswith("mailto:") or not href or href.startswith("#"):
            continue
        result.links.append(LinkItem(text=text, url=href))

    # Leaf-block paragraph extraction — each <p>/<div>/<li>/<h*>/<td>/<blockquote>
    # that does NOT contain another block is treated as ONE paragraph.
    blocks = _leaf_blocks(soup)
    _split_sections(result, blocks)
    result.raw_text = "\n\n".join(blocks)


_BLOCK_TAGS = {"p", "div", "li", "h1", "h2", "h3", "h4", "h5", "h6", "blockquote", "td", "th", "article", "section", "header", "footer"}


def _leaf_blocks(soup) -> List[str]:
    """Return a list of paragraph-level text blocks.

    A "leaf block" is any block-level tag that does NOT contain another
    block-level tag. Its inner text (with <br> converted to a space) becomes
    ONE paragraph. This preserves original DOCX-style paragraph groupings and
    prevents 47-line explosions from every soft line break.
    """
    seen = set()
    blocks: List[str] = []

    for tag in soup.find_all(list(_BLOCK_TAGS)):
        # only leaf blocks
        if any(child.name in _BLOCK_TAGS for child in tag.find_all(True, recursive=True)):
            continue
        # convert <br> to space so wrapped lines stay in one paragraph
        for br in tag.find_all("br"):
            br.replace_with(" ")
        txt = normalize_ws(tag.get_text(" ", strip=True))
        if not txt:
            continue
        key = txt.lower()
        if key in seen:
            continue
        seen.add(key)
        blocks.append(txt)

    # Fallback for emails that use no block markup
    if not blocks:
        raw = soup.get_text("\n")
        for chunk in re.split(r"\n\s*\n", raw):
            t = normalize_ws(chunk)
            if t:
                blocks.append(t)
    return blocks


def _parse_plain_body(text: str, result: ParsedEmail):
    lines = [normalize_ws(l) for l in text.split("\n") if normalize_ws(l)]
    url_re = re.compile(r"https?://[^\s)>\]]+")
    for l in lines:
        for m in url_re.findall(l):
            result.links.append(LinkItem(text=l, url=m))
    _split_sections(result, lines)
    result.raw_text = "\n".join(lines)


# ---------- Section splitter ----------

def _split_sections(result: ParsedEmail, lines: List[str]):
    if not lines:
        return

    greet_idx = None
    for i, l in enumerate(lines[:6]):
        if is_greeting(l):
            greet_idx = i
            break

    if greet_idx is not None:
        result.greeting = lines[greet_idx]
        rest = lines[greet_idx + 1:]
    else:
        result.greeting = ""
        rest = list(lines)

    if not result.subject:
        for l in lines[:6]:
            m = re.match(r"^\s*subject\s*[:\-]\s*(.+)$", l, re.I)
            if m:
                result.subject = normalize_ws(m.group(1))
                rest = [x for x in rest if x != l]
                break

    footer_lines: List[str] = []
    body_lines = list(rest)
    # collect any trailing sequence containing footer keywords
    while body_lines and is_footer_line(body_lines[-1]):
        footer_lines.insert(0, body_lines.pop())
    result.footer = "\n".join(footer_lines).strip()

    if not result.cta:
        for link in result.links:
            if is_cta_text(link.text):
                result.cta = LinkItem(text=link.text, url=link.url)
                break
    if not result.cta:
        for l in body_lines:
            if is_cta_text(l):
                result.cta = LinkItem(text=l, url="")
                break

    cta_text = (result.cta.text if result.cta else "").strip().lower()
    filtered = []
    for l in body_lines:
        if cta_text and l.strip().lower() == cta_text:
            continue
        # skip lines equal to a link's text (avoid CTA duplication when link text is also a paragraph)
        if any(l.strip().lower() == (lk.text or "").strip().lower() for lk in result.links if is_cta_text(lk.text)):
            continue
        filtered.append(l)
    result.body_paragraphs = filtered


# ---------- Comparator ----------

def _status_from_pct(pct: float) -> str:
    if pct >= 0.98:
        return "match"
    if pct >= 0.80:
        return "warning"
    return "mismatch"


def compare_field(a: str, b: str, glossary, mode: str) -> Dict:
    if not a and not b:
        return {"mockup": "", "output": "", "similarity": 1.0, "status": "match"}
    sim = similarity(a, b, glossary, mode)
    return {"mockup": a or "", "output": b or "", "similarity": round(sim, 4), "status": _status_from_pct(sim)}


def _word_diff(a: str, b: str, glossary, mode: str):
    """Return two lists of {text, status} tokens for inline highlighting.

    status ∈ {"match", "del", "add"}
      del → present only in mockup (a)
      add → present only in output (b)
    Comparison is done on **normalized** words but display keeps original casing.
    """
    a_words_disp = a.split()
    b_words_disp = b.split()

    a_norm_tokens = norm_for_compare(a, glossary, mode).split()
    b_norm_tokens = norm_for_compare(b, glossary, mode).split()

    # If normalization changed word count we can't align 1:1 → fallback to raw.
    if len(a_norm_tokens) != len(a_words_disp) or len(b_norm_tokens) != len(b_words_disp):
        a_norm_tokens = [w.lower() for w in a_words_disp]
        b_norm_tokens = [w.lower() for w in b_words_disp]

    sm = SequenceMatcher(None, a_norm_tokens, b_norm_tokens)
    a_out, b_out = [], []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for k in range(i2 - i1):
                a_out.append({"text": a_words_disp[i1 + k], "status": "match"})
                b_out.append({"text": b_words_disp[j1 + k], "status": "match"})
        elif tag == "delete":
            for k in range(i1, i2):
                a_out.append({"text": a_words_disp[k], "status": "del"})
        elif tag == "insert":
            for k in range(j1, j2):
                b_out.append({"text": b_words_disp[k], "status": "add"})
        elif tag == "replace":
            for k in range(i1, i2):
                a_out.append({"text": a_words_disp[k], "status": "del"})
            for k in range(j1, j2):
                b_out.append({"text": b_words_disp[k], "status": "add"})
    return a_out, b_out


def compare_paragraphs(a: List[str], b: List[str], glossary, mode: str) -> Tuple[List[Dict], float]:
    """Best-match paragraph alignment with **placement checking** + word-level diff.

    Each row includes:
      - text_match: bool  (content matches, ignoring position)
      - placement: "correct" | "incorrect" | "missing"
      - mockup_index / output_index
      - mockup_words / output_words  (inline word-level statuses)
    """
    a_norm = [norm_for_compare(x, glossary, mode) for x in a]
    b_norm = [norm_for_compare(x, glossary, mode) for x in b]
    used_b: Dict[int, int] = {}  # b_index -> a_index (owner)
    rows: List[Dict] = []
    sims: List[float] = []

    # First pass: greedy best-match for every mockup paragraph
    for i, ai in enumerate(a):
        best_j, best_sim = -1, 0.0
        for j, bj in enumerate(b_norm):
            if j in used_b:
                continue
            s = SequenceMatcher(None, a_norm[i], bj).ratio() if (a_norm[i] or bj) else 1.0
            if s > best_sim:
                best_sim, best_j = s, j
        if best_j >= 0 and best_sim >= 0.5:
            used_b[best_j] = i
            text_match = best_sim >= 0.98
            placement = "correct" if best_j == i else "incorrect"
            if text_match and placement == "correct":
                status = "match"
            elif text_match and placement == "incorrect":
                status = "warning"  # content matches but placement is wrong
            elif best_sim >= 0.80:
                status = "warning"
            else:
                status = "mismatch"
            aw, bw = _word_diff(a[i], b[best_j], glossary, mode)
            rows.append({
                "mockup": a[i], "output": b[best_j],
                "mockup_index": i, "output_index": best_j,
                "similarity": round(best_sim, 4),
                "text_match": text_match,
                "placement": placement,
                "status": status,
                "mockup_words": aw,
                "output_words": bw,
            })
            sims.append(best_sim)
        else:
            aw, _ = _word_diff(a[i], "", glossary, mode)
            rows.append({
                "mockup": a[i], "output": "",
                "mockup_index": i, "output_index": -1,
                "similarity": 0.0,
                "text_match": False,
                "placement": "missing",
                "status": "mismatch",
                "mockup_words": aw,
                "output_words": [],
            })
            sims.append(0.0)

    # Extras from output — content only in output, not in mockup
    for j, bj in enumerate(b):
        if j not in used_b:
            _, bw = _word_diff("", b[j], glossary, mode)
            rows.append({
                "mockup": "", "output": b[j],
                "mockup_index": -1, "output_index": j,
                "similarity": 0.0,
                "text_match": False,
                "placement": "missing",
                "status": "mismatch",
                "mockup_words": [],
                "output_words": bw,
            })
            sims.append(0.0)

    avg = sum(sims) / len(sims) if sims else 1.0
    return rows, round(avg, 4)


def _url_norm(u: str) -> str:
    return (u or "").strip().rstrip("/").lower()


def compare_links(a: List[LinkItem], b: List[LinkItem], glossary, mode: str) -> Tuple[List[Dict], Dict]:
    used_b = set()
    rows: List[Dict] = []
    text_hits = url_hits = both_hits = 0
    total = max(len(a), len(b)) or 1

    for la in a:
        best_j, best_score = -1, -1.0
        for j, lb in enumerate(b):
            if j in used_b:
                continue
            t_sim = SequenceMatcher(None, norm_for_compare(la.text, glossary, mode), norm_for_compare(lb.text, glossary, mode)).ratio()
            u_sim = SequenceMatcher(None, _url_norm(la.url), _url_norm(lb.url)).ratio()
            score = (t_sim + u_sim) / 2
            if score > best_score:
                best_score, best_j = score, j
        if best_j >= 0 and best_score >= 0.5:
            used_b.add(best_j)
            t_match = norm_for_compare(la.text, glossary, mode) == norm_for_compare(b[best_j].text, glossary, mode)
            u_match = _url_norm(la.url) == _url_norm(b[best_j].url)
            if t_match: text_hits += 1
            if u_match: url_hits += 1
            if t_match and u_match: both_hits += 1
            status = "match" if (t_match and u_match) else ("warning" if (t_match or u_match) else "mismatch")
            rows.append({
                "mockup_text": la.text, "mockup_url": la.url,
                "output_text": b[best_j].text, "output_url": b[best_j].url,
                "text_match": t_match, "url_match": u_match, "status": status
            })
        else:
            rows.append({"mockup_text": la.text, "mockup_url": la.url,
                         "output_text": "", "output_url": "",
                         "text_match": False, "url_match": False, "status": "mismatch"})
    for j, lb in enumerate(b):
        if j not in used_b:
            rows.append({"mockup_text": "", "mockup_url": "",
                         "output_text": lb.text, "output_url": lb.url,
                         "text_match": False, "url_match": False, "status": "mismatch"})
    stats = {
        "matched": both_hits,
        "text_matched": text_hits,
        "url_matched": url_hits,
        "total": total,
        "pct": round(both_hits / total, 4) if total else 1.0,
    }
    return rows, stats


def compare_images(a: List[ImageItem], b: List[ImageItem], glossary, mode: str) -> Tuple[List[Dict], Dict]:
    used_b = set()
    rows: List[Dict] = []
    total = max(len(a), len(b)) or 1
    matched = 0

    for ia in a:
        best_j, best_score = -1, -1.0
        for j, ib in enumerate(b):
            if j in used_b:
                continue
            alt_sim = SequenceMatcher(None, norm_for_compare(ia.alt, glossary, mode), norm_for_compare(ib.alt, glossary, mode)).ratio()
            name_sim = SequenceMatcher(None, ia.filename.lower(), ib.filename.lower()).ratio()
            score = max(alt_sim, name_sim)
            if score > best_score:
                best_score, best_j = score, j
        if best_j >= 0 and best_score >= 0.4:
            used_b.add(best_j)
            ib = b[best_j]
            alt_match = bool(ia.alt or ib.alt) and norm_for_compare(ia.alt, glossary, mode) == norm_for_compare(ib.alt, glossary, mode)
            name_match = ia.filename.lower() == ib.filename.lower()
            ok = alt_match or name_match
            if ok:
                matched += 1
            rendered_alt = ib.alt if ib.alt else ("no rendered image" if not ib.filename else ib.filename)
            rows.append({
                "mockup_filename": ia.filename, "mockup_alt": ia.alt,
                "output_filename": ib.filename, "output_alt": ib.alt,
                "rendered": True,
                "rendered_label": rendered_alt,
                "alt_match": alt_match, "name_match": name_match,
                "status": "match" if ok else "warning"
            })
        else:
            # image in mockup but not rendered in output
            rows.append({
                "mockup_filename": ia.filename, "mockup_alt": ia.alt,
                "output_filename": "", "output_alt": "",
                "rendered": False,
                "rendered_label": "no rendered image",
                "alt_match": False, "name_match": False,
                "status": "mismatch"
            })
    for j, ib in enumerate(b):
        if j not in used_b:
            rows.append({
                "mockup_filename": "", "mockup_alt": "",
                "output_filename": ib.filename, "output_alt": ib.alt,
                "rendered": True,
                "rendered_label": ib.alt or ib.filename or "no rendered image",
                "alt_match": False, "name_match": False,
                "status": "mismatch"
            })
    stats = {"matched": matched, "total": total, "pct": round(matched / total, 4) if total else 1.0}
    return rows, stats


def compare_cta(a: Optional[LinkItem], b: Optional[LinkItem], glossary, mode: str) -> Dict:
    a = a or LinkItem("", "")
    b = b or LinkItem("", "")
    text_res = compare_field(a.text, b.text, glossary, mode)
    url_match = _url_norm(a.url) == _url_norm(b.url)
    if not (a.text or b.text or a.url or b.url):
        pct = 1.0
        status = "match"  # no CTA in either — do not penalise
    else:
        pct = (text_res["similarity"] + (1.0 if url_match else 0.0)) / 2
        if text_res["status"] == "match" and url_match:
            status = "match"
        elif text_res["status"] == "mismatch" and not url_match:
            status = "mismatch"
        else:
            status = "warning"
    return {
        "mockup_text": a.text, "mockup_url": a.url,
        "output_text": b.text, "output_url": b.url,
        "text_similarity": text_res["similarity"], "url_match": url_match,
        "similarity": round(pct, 4),
        "status": status,
    }


# Weighted overall
WEIGHTS = {
    "subject": 0.15,
    "greeting": 0.05,
    "body": 0.40,
    "cta": 0.15,
    "links": 0.10,
    "images": 0.10,
    "footer": 0.05,
}


def compare(mockup: ParsedEmail, output: ParsedEmail,
            glossary: List[Dict[str, str]] | None = None,
            mode: str = "smart", use_glossary: bool = True) -> Dict:
    glossary = (glossary or []) if use_glossary else []
    mode = mode if mode in ("smart", "strict") else "smart"

    subj = compare_field(mockup.subject, output.subject, glossary, mode)
    greet = compare_field(mockup.greeting, output.greeting, glossary, mode)
    cta = compare_cta(mockup.cta, output.cta, glossary, mode)
    footer = compare_field(mockup.footer, output.footer, glossary, mode)
    body_rows, body_pct = compare_paragraphs(mockup.body_paragraphs, output.body_paragraphs, glossary, mode)
    link_rows, link_stats = compare_links(mockup.links, output.links, glossary, mode)
    img_rows, img_stats = compare_images(mockup.images, output.images, glossary, mode)

    scores = {
        "subject": round(subj["similarity"] * 100, 2),
        "greeting": round(greet["similarity"] * 100, 2),
        "body": round(body_pct * 100, 2),
        "cta": round(cta["similarity"] * 100, 2),
        "links": round(link_stats["pct"] * 100, 2),
        "images": round(img_stats["pct"] * 100, 2),
        "footer": round(footer["similarity"] * 100, 2),
    }
    overall_score = round(sum(scores[k] * WEIGHTS[k] for k in scores), 2)

    total_checks = 4 + len(body_rows) + len(link_rows) + len(img_rows)
    all_states = [subj["status"], greet["status"], cta["status"], footer["status"]] + \
                 [r["status"] for r in body_rows] + \
                 [r["status"] for r in link_rows] + \
                 [r["status"] for r in img_rows]
    matches = sum(1 for s in all_states if s == "match")
    warnings = sum(1 for s in all_states if s == "warning")
    mismatches = sum(1 for s in all_states if s == "mismatch")

    if overall_score >= 98:
        overall_status = "match"
    elif overall_score >= 80:
        overall_status = "warning"
    else:
        overall_status = "mismatch"

    return {
        "mode": mode,
        "use_glossary": use_glossary,
        "subject": subj,
        "greeting": greet,
        "cta": cta,
        "footer": footer,
        "body": body_rows,
        "links": link_rows,
        "images": img_rows,
        "link_stats": link_stats,
        "image_stats": img_stats,
        "scores": scores,
        "summary": {
            "overall_score": overall_score,
            "overall": overall_status,
            "total_checks": total_checks,
            "match": matches,
            "warning": warnings,
            "mismatch": mismatches,
            "score_percent": overall_score,  # backward compat
        },
        "mockup_parsed": mockup.to_dict(),
        "output_parsed": output.to_dict(),
    }


# ---------- Dispatchers ----------

def parse_output(filename: str, file_bytes: bytes) -> ParsedEmail:
    lower = (filename or "").lower()
    if lower.endswith(".eml"):
        return parse_eml(file_bytes)
    if lower.endswith(".msg"):
        return parse_msg(file_bytes)
    raise ValueError(f"Unsupported output file type: {filename}")


def parse_mockup(filename: str, file_bytes: bytes) -> ParsedEmail:
    lower = (filename or "").lower()
    if lower.endswith(".docx"):
        return parse_docx(file_bytes)
    raise ValueError(f"Unsupported mockup file type: {filename}")
