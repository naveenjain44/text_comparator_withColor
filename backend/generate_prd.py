"""Generate PRD as .docx (offline, python-docx). Run once at startup or manually."""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUT_PATH = Path(__file__).parent.parent / "docs" / "UAT_Text_Comparator_PRD.docx"


def _h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x11, 0x11, 0x11)


def _p(doc, text, bold=False, mono=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if mono:
        r.font.name = "Consolas"
        r.font.size = Pt(10)
    else:
        r.font.size = Pt(11)
    return p


def _code(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(code_text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    return p


def _bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def _number(doc, text):
    doc.add_paragraph(text, style="List Number")


def build():
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run("UAT Text Comparator")
    r.bold = True
    r.font.size = Pt(26)

    sub = doc.add_paragraph()
    r = sub.add_run("Product Requirements Document · Offline localhost tool for email mockup UAT")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # Overview
    _h(doc, "1. Overview", level=1)
    _p(doc,
       "UAT Text Comparator is a Python + React tool that runs entirely on your local machine "
       "(no internet, no API keys) to perform apple-to-apple UAT between a marketing email mockup "
       "(.docx) and its rendered output (.eml or Outlook .msg). It produces PDF, HTML, JSON, and "
       "CSV artifacts, keeps a local run history, and lets teams define allowed variations via a "
       "glossary/case-rule editor."
       )

    # Goals
    _h(doc, "2. Goals & Non-Goals", level=1)
    _p(doc, "Goals", bold=True)
    _bullet(doc, "Compare subject, greeting, CTA (text+URL), body paragraphs, footer, hyperlinks, and images (filename + alt).")
    _bullet(doc, "Support batch mode via folder or ZIP with auto-pairing by filename stem.")
    _bullet(doc, "Export PDF, HTML, CSV, and JSON reports. Persist run history on disk (SQLite).")
    _bullet(doc, "Editable glossary of allowed variations (case-rules), persisted as JSON.")
    _bullet(doc, "Ignore strikethrough runs in the DOCX mockup so scores reflect current copy only.")
    _bullet(doc, "Work OFFLINE — verified with internet disconnected.")
    _p(doc, "Non-Goals", bold=True)
    _bullet(doc, "Live URL reachability checks (would require internet — deliberately excluded).")
    _bullet(doc, "Rendering fidelity (pixel-level image diff of the rendered email).")
    _bullet(doc, "Cloud-hosted collaboration — this is a local team tool distributed via Git.")

    # Personas
    _h(doc, "3. Personas", level=1)
    _bullet(doc, "QA / UAT engineer verifying email campaign builds before release.")
    _bullet(doc, "Marketing manager checking that approved copy is what actually ships.")
    _bullet(doc, "Dev-team lead integrating the tool into a repeatable UAT workflow via GitHub.")

    # Features
    _h(doc, "4. Features", level=1)
    _p(doc, "4.1 Single comparison", bold=True)
    _bullet(doc, "Upload mockup (.docx) and output (.eml / .msg).")
    _bullet(doc, "Toggle STRICT (exact) vs SMART (glossary + case-insensitive + punctuation-tolerant).")
    _bullet(doc, "Score cards: Overall (weighted), Subject, Body, URLs, Images.")
    _bullet(doc, "Sub-tabs: Subject, Body Diff, CTA, URLs, Images, Footer.")
    _bullet(doc, "Export PDF and HTML from the run.")

    _p(doc, "4.2 Batch mode", bold=True)
    _bullet(doc, "Drop many pairs or a .zip; pairs auto-match by filename stem (e.g. campaign1.docx ↔ campaign1.eml).")
    _bullet(doc, "Per-pair report + CSV summary.")

    _p(doc, "4.3 Glossary / Case Rules", bold=True)
    _bullet(doc, 'Editable rules like "Hi" ↔ "Hey" persisted in backend/glossary.json.')
    _bullet(doc, "Rules are applied in SMART mode; disabled in STRICT mode.")

    _p(doc, "4.4 History", bold=True)
    _bullet(doc, "Every comparison is written to a local SQLite DB (backend/data/history.db).")
    _bullet(doc, "Right-side history panel lists last 50 runs with score, filenames, timestamp.")
    _bullet(doc, "Any past run can be re-downloaded as PDF or HTML.")

    _p(doc, "4.5 Strikethrough handling", bold=True)
    _bullet(doc, "Text formatted with w:strike or w:dstrike in the DOCX is fully excluded from comparison.")
    _bullet(doc, "Applies to plain runs and to text inside hyperlinks.")

    # Tech stack
    _h(doc, "5. Tech Stack", level=1)
    _bullet(doc, "Backend: Python 3.10+, FastAPI, Uvicorn, python-docx, extract-msg, BeautifulSoup4 + lxml, ReportLab, SQLite (stdlib).")
    _bullet(doc, "Frontend: React 19 (CRA + CRACO), TailwindCSS, shadcn/ui, lucide-react, sonner (toasts).")
    _bullet(doc, "Persistence: local JSON (glossary), local SQLite (history). No cloud dependencies.")
    _bullet(doc, "No LLM, no external HTTP calls at runtime.")

    # Folder tree
    _h(doc, "6. Code Folder Structure", level=1)
    _code(doc,
"""uat-text-comparator/
├── backend/
│   ├── server.py              FastAPI app + endpoints
│   ├── comparator.py          DOCX / EML / MSG parsers + diff logic
│   ├── pdf_report.py          ReportLab PDF builder
│   ├── html_report.py         Self-contained HTML report builder
│   ├── history.py             SQLite history store
│   ├── generate_prd.py        Regenerates this PRD.docx
│   ├── glossary.json          Editable case rules
│   ├── data/
│   │   └── history.db         SQLite (auto-created)
│   ├── requirements.txt
│   └── .env                   (MONGO_URL/DB_NAME kept but unused; safe to remove)
├── frontend/
│   ├── src/
│   │   ├── App.js
│   │   ├── pages/
│   │   │   ├── Compare.jsx    Main comparison UI + history sidebar
│   │   │   ├── Batch.jsx      Batch mode + CSV export
│   │   │   ├── Glossary.jsx   Case-rule editor
│   │   │   └── Guide.jsx      In-app setup guide
│   │   ├── components/
│   │   │   ├── Layout.jsx     Header + tabs
│   │   │   ├── UploadZone.jsx
│   │   │   ├── DiffReport.jsx
│   │   │   └── HistorySidebar.jsx
│   │   └── lib/api.js         Auto-detects localhost vs. hosted
│   ├── .env                   REACT_APP_BACKEND_URL (localhost by default)
│   └── package.json
├── docs/
│   └── UAT_Text_Comparator_PRD.docx  (this file)
└── README.md
""")

    # Install
    _h(doc, "7. Installation — macOS", level=1)
    _number(doc, "Install Python 3.10+ and Node 18+ (via Homebrew: brew install python@3.11 node).")
    _number(doc, "Clone your repo and enter the folder.")
    _code(doc, "git clone https://github.com/<your-user>/uat-text-comparator.git\ncd uat-text-comparator")
    _number(doc, "Backend setup:")
    _code(doc, "cd backend\npython3 -m venv .venv\nsource .venv/bin/activate\npip install -r requirements.txt\nuvicorn server:app --host 127.0.0.1 --port 8001")
    _number(doc, "Frontend setup (new terminal):")
    _code(doc, "cd frontend\nyarn install\nyarn start   # opens http://localhost:3000")
    _number(doc, "Open http://localhost:3000 in your browser. Disconnect the internet — the tool keeps working.")

    _h(doc, "8. Installation — Windows", level=1)
    _number(doc, "Install Python 3.10+ from python.org and Node.js 18+ LTS from nodejs.org.")
    _number(doc, "Install Yarn: npm install -g yarn (one-time, needs internet once).")
    _number(doc, "Clone the repo (or download the ZIP) and open a PowerShell terminal in the folder.")
    _number(doc, "Backend setup:")
    _code(doc, "cd backend\npython -m venv .venv\n.venv\\Scripts\\activate\npip install -r requirements.txt\nuvicorn server:app --host 127.0.0.1 --port 8001")
    _number(doc, "Frontend setup (new PowerShell):")
    _code(doc, "cd frontend\nyarn install\nyarn start")
    _number(doc, "Open http://localhost:3000 in Edge / Chrome. Everything runs locally.")

    _h(doc, "9. Offline Guarantee", level=1)
    _bullet(doc, "The frontend auto-detects hostname == 'localhost' and points to http://localhost:8001 (no env change needed).")
    _bullet(doc, "All parsing, PDF, and HTML generation happen inside Python — no cloud services.")
    _bullet(doc, "Glossary is a JSON file on disk. History is a local SQLite DB.")
    _bullet(doc, "You can unplug the network cable after yarn install / pip install — everything keeps working.")

    _h(doc, "10. Scoring Model", level=1)
    _p(doc, "Each section produces a similarity percentage. The overall score is a weighted average:")
    _bullet(doc, "Subject 15% · Greeting 5% · Body 40% · CTA 15% · Links 10% · Images 10% · Footer 5%.")
    _bullet(doc, "Body uses best-match paragraph alignment via difflib SequenceMatcher.")
    _bullet(doc, "Links & Images use matched pairs over total set to derive percentage.")
    _bullet(doc, "Status thresholds: MATCH ≥ 98%, WARNING 80-97%, MISMATCH < 80%.")

    _h(doc, "11. API (for CI / advanced users)", level=1)
    _bullet(doc, "POST /api/compare  (multipart: mockup, output, mode, use_glossary)  →  full report JSON")
    _bullet(doc, "POST /api/batch    (multipart: files[])  →  list of reports + csv_rows")
    _bullet(doc, "POST /api/export/pdf  ({report, mockup_filename, output_filename})  →  PDF")
    _bullet(doc, "POST /api/export/html ({report, mockup_filename, output_filename})  →  HTML")
    _bullet(doc, "POST /api/export/csv  ({rows})  →  CSV")
    _bullet(doc, "GET  /api/history · /api/history/{id} · /api/history/{id}/pdf · /api/history/{id}/html")
    _bullet(doc, "GET/POST /api/glossary  →  read / persist rules")

    _h(doc, "12. Roadmap", level=1)
    _bullet(doc, "Inline image thumbnail preview inside the diff.")
    _bullet(doc, "Word-level highlight inside a paragraph diff.")
    _bullet(doc, "One-click PyInstaller bundle (desktop .exe / .app).")
    _bullet(doc, "Multi-project profiles: per-brand section heuristics + weights.")

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    p = build()
    print(f"Wrote {p}")
